from pyzabbix import ZabbixAPI
from datetime import datetime
from tempfile import TemporaryDirectory
from docx import Document
import re, time, pandas, requests
import dateutil.parser

class ZabbixReport(object):
    """
    ZabbixAPIを使用して各種のデータを収集整形する
    """

    def __init__(self, server = "http://127.0.0.1/zabbix"):
        """
        初期設定を行う
        デフォルトでサーバのURLはローカルのzabbixサーバになっている
        """
        self.server = server
        self.zapi = ZabbixAPI(server = self.server)
        self.hosts_dictionary = {}

    def update_hosts_dictionary(self):
        """
        ホストの辞書データを更新する
        辞書は self.hosts_dictionary = {"ホスト":"ホストのID", ...}となっている
        """
        for host in self.zapi.host.get(output=["host"]):
            self.hosts_dictionary[host["host"]] = host["hostid"]

    def items_dictionary_of_host(self, hostid):
        """
        特定のホストに属するアイテムの辞書を得る
        この関数はホストのhostidを必要とする
        辞書は {"アイテムのkey":"アイテムID",....}となっている
        """
        items_dictionary_of_host = {}
        items_list = self.zapi.item.get(hostids=hostid, output=["key_"])

        for item in items_list:
            items_dictionary_of_host[item["key_"]] = item["itemid"]
        return items_dictionary_of_host

    def attribute_of_host(self, hostid):
        return self.zapi.host.get(hostids = hostid)[0]

    def string_to_Unix_time(self, time_string):
        """
        文字列表現の時間をUNIX時間に変換する
        """
        now = datetime.now()
        if time_string == 'last_month':
            time = datetime(now.year, now.month -1, 1)
        elif time_string == 'this_month':
            time = datetime(now.year, now.month, 1)
        else:
            time = dateutil.parser.parse(time_string)
        Unix_time = int(time.timestamp())
        return Unix_time

    def Unix_time_to_string(self, Unix_time):
        return datetime.fromtimestamp(int(Unix_time)).isoformat(' ')

    def history_of_item(self, itemid, time_from, time_till):
        time_from = self.string_to_Unix_time(time_from)
        time_till = self.string_to_Unix_time(time_till)
        value_type = self.zapi.item.get(itemids =itemid, output=['value_type'])[0]['value_type']
        return self.zapi.history.get(itemids =itemid, history=value_type, time_from = time_from, time_till = time_till)

    def save_history_as_csv(self, itemid, time_from, time_till, name_saving_file = "hoge.csv"):
        """
        特定のアイテムのヒストリーをcsvファイルに保存する。
        ヒストリの時間はUNIXタイムスタンプからISOフォーマットに変換する。
        """
        item_attr = self.zapi.item.get(itemid = itemid, output=['key_', 'name', 'hostid'])[0]
        host_attr = self.zapi.host.get(hostid = item_attr['hostid'], output=['name'])[0]
        item_info = 'item_info -->'
        item_info = item_info + str(item_attr)
        host_info = 'host_info -->'
        host_info = host_info + str(host_info)
        history_of_item = self.history_of_item(itemid = itemid, time_from = time_from, time_till = time_till)
        for row in history_of_item:
            row['clock'] = self.Unix_time_to_string(row['clock'])
        csv_file = open(name_saving_file, mode ='w')
        csv_file.write(host_info + '\n' + item_info)
        csv_file.close()
        history_dataframe = pandas.DataFrame(history_of_item)
        history_dataframe['clock', 'value'].to_csv(name_saving_file, mode = 'a')
        return name_saving_file

    def save_graph_image(self, graphid, time_from, time_till, width=1600, height=900, save_as='graph.png'):
        """
        グラフの画像を保存する。
        グラフのID、グラフにする期間の始まりと終わりを必要とする。
        オプションで、グラフの縦・横のサイズ、画像ファイルの名前を指定できる。
        返す値は保存したグラフ画像のパス。
        """
        period = self.string_to_Unix_time(time_till) - self.string_to_Unix_time(time_from)
        time_from_as_datetime = dateutil.parser.parse(time_from)
        parameters = {'graphid':graphid, 'width':width, 'height':height, 'stime':time_from_as_datetime.strftime("%Y%m%d%H%M%S"), 'period':period }
        response = requests.get(self.server + "/chart2.php", params = parameters, cookies = {'zbx_sessionid':self.zapi.auth},)
        file = open(save_as, 'wb')
        file.write(response.content)
        file.close()
        return save_as

    def save_graph_images_with(self, hostid, search_word, time_from, time_till):
        """
        特定ホストのグラフについて、名前にsearch_wordを含むものをすべて保存する。
        返す値は、保存したグラフ画像の相対パスのリスト。
        """
        saved_graph_pathes = []
        for graph in self.zapi.graph.get(hostids = [host_id], search = {'name':[search_word]}, output = ['name']):
            graph_name = graph['name']
            graph_id = graph['graphid']
            saved_graph_pathes.appned(self.save_graph_images(graphid = graph_id, time_from = time_from, time_till = time_till, save_as = graph_name))
        return saved_graph_pathes

    def make_report_for_kk(self, kk_name, time_from = 'last_month', time_till = 'this_month'):
        """
        特定の加入機関のレポートを作成する。
        kk_nameが加入機関の名前、レポートにする期間についても指定可能（デフォルトで先月の一ヶ月間）
        """
        time_from = self.string_to_Unix_time(time_from)
        time_till = self.string_to_Unix_time(time_till)
        hosts_of_kk = {}
        saved_graph_names = []
        for host in self.zapi.host.get(output=['host', 'name'], search = {'name':kk_name}):
            hosts_of_kk[host['name']] = host['hostid']

        for host_name, host_id in hosts_of_kk.items():
            graph_images = []
            if 'kk' in host_name:
                graph_images.extend(self.save_graph_images_with(hostid = host_id, search_word = 'traffic', time_from = time_from, time_till = time_till))
            graph_images.extend(self.save_graph_images_with(hostid = host_id, search_word = 'optpower', time_from = time_from, time_till = time_till)) 
        graph_images.reverse()
        while graph_images:
            paste_to_paper(graph_images.pop()) #レポートのフォーマットに画像を貼り付けていく処理
        return report_name

    def make_report_from_screen(self, screenid, time_from = 'last_month',  time_till = 'this_month', template='', save_as = 'ReportFromScreen.docx'):
        """
        make report in Word file format from screen.
        """
        if template:
            doc = Document(template)
        else:
            doc = Document()
        doc.add_heading(self.zapi.screen.get(screenids = screenid, output = ['name'])[0]['name'], 0)

        tmp_dir = TemporaryDirectory(prefix = 'screen' + screenid)

        screenitems = pandas.DataFrame(
            self.zapi.screenitem.get(
                screenids = screenid,
                output = ['x','y', 'resourceid', 'height', 'width', 'resourcetype']
            )
        )
        for index in screenitems.sort(columns=['y', 'x']).index:
            screenitem = screenitems.loc[index]
            if screenitem['resourcetype'] == '0':
                saved_graph = self.save_graph_image(
                    graphid = screenitem['resourceid'],
                    width = screenitem['width'],
                    height = screenitem['height'],
                    time_from = time_from,
                    time_till = time_till,
                    save_as = tmp_dir.name + '/' + screenitem['resourcetype'] + screenitem['resourceid'] + '.png'
                )
                doc.add_picture(saved_graph,
                                width = doc.sections[0].page_width - doc.sections[0].right_margin - doc.sections[0].left_margin)
        doc.save(save_as)
        tmp_dir.cleanup()
        del tmp_dir
        return save_as

class Server(object):
    """
    Webサーバを提供するクラス
    リクエストを受け取って、レポートを作成しダウンロードできる  
    """
