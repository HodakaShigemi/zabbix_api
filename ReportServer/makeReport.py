#!/usr/bin/python3
# -*- coding: utf-8

from pyzabbix import ZabbixAPI
from datetime import datetime
from tempfile import TemporaryDirectory
from docx import Document
import re, time, pandas, requests
import dateutil.parser

class ZabbixReportAPI(object):
    """
    ZabbixAPIを使用して各種のデータを収集整形する
    """

    def __init__(self, server = "http://127.0.0.1/zabbix", user = 'admin', password ='zabbix'):
        """
        初期設定を行う
        デフォルトでサーバのURLはローカルのzabbixサーバになっている
        """
        self.server = server
        self.zapi = ZabbixAPI(server = self.server)
        self.zapi.login(user = user, password = password)
        self.screens_dictionary = {}
        self.update_screens_dictionary()
        self.hosts_dictionary = {}
        self.update_hosts_dictionary()

    def update_screens_dictionary(self):
        """
        スクリーンの辞書データを更新する
        辞書は self.screens_dictionary = {"スクリーン名":"スクリーンID", ...}となっている
        """
        for screen in self.zapi.screen.get(output=["name"]):
            self.screens_dictionary[screen["name"]] = screen["screenid"]

    def update_hosts_dictionary(self):
        """
        ホストの辞書データを更新する
        辞書は self.hosts_dictionary = {"表示名":"ホストのID", ...}となっている
        """
        for host in self.zapi.host.get(output=["name"]):
            self.hosts_dictionary[host["name"]] = host["hostid"]

    def items_dictionary_of_host(self, hostid):
        """
        特定のホストに属するアイテムの辞書を得る
        この関数はホストのhostidを必要とする
        辞書は {"アイテムのkey":"アイテムID",....}となっている
        """
        items_dictionary_of_host = {}
        items_list = self.zapi.item.get(hostids=hostid, output=["key_"])

        for item in items_list:
            items_dictionary_of_host[item["key_"]] = item["itemid"]
        return items_dictionary_of_host

    def to_Unix_time(self, time_val):
        """
        渡された引数をUNIX時間(整数値)に変換して返す
        """
        now = datetime.now()
        if type(time_val) == (int or float):
            return int(time_val)
        if type(time_val) == datetime:
            time = time_val
        elif time_val == 'last_month':
            time = datetime(now.year, now.month -1, 1)
        elif time_val == 'this_month':
            time = datetime(now.year, now.month, 1)
        else:
            time = dateutil.parser.parse(time_val)
        Unix_time = int(time.timestamp())
        return Unix_time

    def Unix_time_to_string(self, Unix_time):
        return datetime.fromtimestamp(int(Unix_time)).isoformat(' ')

    def history_of_item(self, itemid, time_from, time_till):
        time_from = self.to_Unix_time(time_from)
        time_till = self.to_Unix_time(time_till)
        value_type = self.zapi.item.get(itemids =itemid, output=['value_type'])[0]['value_type']
        return self.zapi.history.get(itemids =itemid, history=value_type, time_from = time_from, time_till = time_till)

    def save_history_as_csv(self, itemid, time_from, time_till, save_as = "hoge.csv"):
        """
        特定のアイテムのヒストリーをcsvファイルに保存する。
        ヒストリの時間はUNIXタイムスタンプからISOフォーマットに変換する。
        """
        item_attr = self.zapi.item.get(itemid = itemid, output=['key_', 'name', 'hostid'])[0]
        host_attr = self.zapi.host.get(hostid = item_attr['hostid'], output=['name'])[0]
        item_info = 'item_info -->,'
        item_info = item_info + str(item_attr)
        host_info = 'host_info -->,'
        host_info = host_info + str(host_attr)
        csv_file = open(save_as, mode ='w')
        csv_file.write(host_info + '\n' + item_info + '\n')
        csv_file.close()
        history_of_item = self.history_of_item(itemid = itemid, time_from = time_from, time_till = time_till)
        for row in history_of_item:
            row['clock'] = self.Unix_time_to_string(row['clock'])
        history_dataframe = pandas.DataFrame(history_of_item)
        history_dataframe[['clock', 'value']].to_csv(save_as, mode = 'a')
        return save_as

    def save_graph_image(self, graphid, time_from, time_till, width=1600, height=900, save_as='graph.png'):
        """
        グラフの画像を保存する。
        グラフのID、グラフにする期間の始まりと終わりを必要とする。
        オプションで、グラフの縦・横のサイズ、画像ファイルの名前を指定できる。
        返す値は保存したグラフ画像のパス。
        """
        period = self.to_Unix_time(time_till) - self.to_Unix_time(time_from)
        time_from = self.to_Unix_time(time_from)
        time_from_as_datetime = datetime.fromtimestamp(time_from)
        parameters = {'graphid':graphid, 'width':width, 'height':height, 'stime':time_from_as_datetime.strftime("%Y%m%d%H%M%S"), 'period':period }
        response = requests.get(self.server + "/chart2.php", params = parameters, cookies = {'zbx_sessionid':self.zapi.auth},)
        file = open(save_as, 'wb')
        file.write(response.content)
        file.close()
        return save_as

    def save_graph_images_with(self, hostid, search_word, time_from, time_till):
        """
        特定ホストのグラフについて、名前にsearch_wordを含むものをすべて保存する。
        返す値は、保存したグラフ画像の相対パスのリスト。
        """
        saved_graph_pathes = []
        for graph in self.zapi.graph.get(hostids = [host_id], search = {'name':[search_word]}, output = ['name']):
            graph_name = graph['name']
            graph_id = graph['graphid']
            saved_graph_pathes.appned(self.save_graph_images(graphid = graph_id, time_from = time_from, time_till = time_till, save_as = graph_name))
        return saved_graph_pathes

    def save_report_from_screen(self, screenid, time_from = 'last_month',  time_till = 'this_month', template = None, save_as = 'ReportFromScreen.docx'):
        """
        Zabbixのスクリーン画面に登録されているグラフをワードファイルに張り付けて保存する。
        現在グラフのみをワードに張り付けている。
        ワードファイル内のグラフ画像は、すべて縦に並べられる。
        引数としては、スクリーンのID、グラフにする時間の始めと終わり、テンプレートにするファイル、
        保存するワードファイルの名前。を引数に取ることができる。
        """
        if template:
            doc = Document(template)
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    if '{title}' in run.text:
                        run.text = run.text.format(title = self.zapi.screen.get(screenids = screenid, output = ['name'])[0]['name'])
        else:
            doc = Document()
            doc.add_heading(self.zapi.screen.get(screenids = screenid, output = ['name'])[0]['name'], 0)

        tmp_dir = TemporaryDirectory(prefix = 'screen' + screenid)

        screenitems = pandas.DataFrame(
            self.zapi.screenitem.get(
                screenids = screenid,
                output = ['x','y', 'resourceid', 'height', 'width', 'resourcetype']
            )
        )
        for index in screenitems.astype(int).sort(columns=['y', 'x']).index:
            screenitem = screenitems.loc[index]
            if screenitem['resourcetype'] == '0':
                saved_graph = self.save_graph_image(
                    graphid = screenitem['resourceid'],
                    width = screenitem['width'],
                    height = screenitem['height'],
                    time_from = time_from,
                    time_till = time_till,
                    save_as = tmp_dir.name + '/' + screenitem['resourcetype'] + screenitem['resourceid'] + '.png'
                )
                doc.add_picture(saved_graph,
                                width = doc.sections[0].page_width - doc.sections[0].right_margin - doc.sections[0].left_margin)
        doc.save(save_as)
        tmp_dir.cleanup()
        del tmp_dir
        return save_as

    def save_reports_from_screens(self, save_dir, time_from = 'last_month', time_till = 'this_month', template = None):
        """
        ZABBIXに登録されているスクリーンすべてからレポートを作成する
        保存するディレクトリのパスを引数に取る
        """
        regexp_dc = re.compile(r'[0-9]{2}(.).+DC')
        if save_dir[-1] != '/':
            save_dir = save_dir + '/'
        screens = self.zapi.screen.get(output=['name'])
        for screen in screens:
            if '共同調達' in screen['name']:
                dc_name = regexp_dc.search(screen['name']).group().replace('DC', '')
                file_name = dc_name + '.docx'
            else:
                file_name = screen['name'].replace(' ', '').replace('　', '') + '.docx'
            self.save_report_from_screen(
                screenid = screen['screenid'],
                time_from = time_from,
                time_till = time_till,
                template = template,
                save_as = save_dir + file_name
            )
